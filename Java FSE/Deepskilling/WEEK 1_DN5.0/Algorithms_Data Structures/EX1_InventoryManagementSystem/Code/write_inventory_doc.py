import sys
import subprocess
import pkgutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt

out = Path(r'd:\Digital_Nuturing5.0\Java FSE\Deepskilling\WEEK 1_DN5.0\Algorithms_Data Structures\EX1_InventoryManagementSystem\Output\Inventory_Management_System.docx')
out.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)
styles['Heading 1'].font.size = Pt(16)
styles['Heading 2'].font.size = Pt(13)
styles['Heading 3'].font.size = Pt(12)

doc.add_heading('Inventory Management System - Code Explanation', level=1)

p = doc.add_paragraph()
p.add_run('This document explains every major part of the Java inventory management system code. It covers the purpose of each file, the classes used, and the role of each method and statement.')

p = doc.add_paragraph()
p.add_run('Project Overview').bold = True
p = doc.add_paragraph('The inventory management system is a simple Java program that allows a user to add, update, delete, and display products in an inventory. It uses object-oriented programming concepts such as classes, objects, constructors, and methods.')

sec = doc.add_heading('1. InventoryTest.java', level=2)
par = doc.add_paragraph(); par.add_run('Purpose').bold = True
par = doc.add_paragraph('This file is the main driver of the program. It contains the main method, which is the starting point of execution when the program runs.')

items = [
    'public class InventoryTest {',
    'This line declares the class named InventoryTest. Every Java program needs at least one class, and this class contains the main method that starts the application.',
    'public static void main(String[] args) {',
    'This is the main method. Java runs this method first when the program starts.',
    'Inventory inventory = new Inventory();',
    'This creates an object of the Inventory class. The object is used to perform inventory operations such as adding, updating, deleting, and displaying products.',
    'inventory.addProduct(new Product1(101, "Laptop", 20, 55000));',
    'This adds a new product to the inventory. The Product1 object stores product information such as product ID, name, quantity, and price.',
    'inventory.addProduct(new Product1(102, "Mouse", 50, 500));',
    'This adds another product to the inventory.',
    'inventory.addProduct(new Product1(103, "Keyboard", 30, 1200));',
    'This adds a third product to the inventory.',
    'System.out.println("\\n===== Inventory =====");',
    'This prints a heading before displaying the inventory list.',
    'inventory.displayInventory();',
    'This calls the displayInventory method to print all current products stored in the inventory.',
    'inventory.updateProduct(102, 75, 650);',
    'This updates the product with ID 102 by changing its quantity to 75 and its price to 650.',
    'inventory.deleteProduct(103);',
    'This removes the product with ID 103 from the inventory.'
]
for text in items:
    doc.add_paragraph(text)

sec = doc.add_heading('2. Inventory.java', level=2)
par = doc.add_paragraph(); par.add_run('Purpose').bold = True
par = doc.add_paragraph('This file manages the inventory data structure and all operations related to products.')

items = [
    'import java.util.HashMap;',
    'This imports the HashMap class from the Java utility package. It is used to store products in a key-value structure.',
    'private HashMap<Integer, Product1> inventory = new HashMap<>();',
    'This creates a HashMap where the key is the product ID and the value is a Product1 object. This allows fast access to each product using its ID.',
    'public void addProduct(Product1 product) {',
    'This method adds a product to the inventory.',
    'inventory.put(product.productId, product);',
    'The product ID is used as the key, and the product object is stored as the value.',
    'System.out.println("Product added successfully.");',
    'This prints a confirmation message after the product is added.',
    'public void updateProduct(int productId, int quantity, double price) {',
    'This method updates an existing product. It takes the product ID, new quantity, and new price as parameters.',
    'Product1 product = inventory.get(productId);',
    'This searches the inventory map for the product using its product ID.',
    'if (product != null) {',
    'This checks whether the product exists in the inventory.',
    'product.quantity = quantity;',
    'This assigns the new quantity to the product.',
    'product.price = price;',
    'This assigns the new price to the product.',
    'System.out.println("Product updated successfully.");',
    'This confirms the update.',
    '} else {',
    'If the product is not found, this part runs.',
    'System.out.println("Product not found.");',
    'This prints an error message.',
    'public void deleteProduct(int productId) {',
    'This method deletes a product from the inventory.',
    'if (inventory.remove(productId) != null) {',
    'This removes the product from the map using its ID and checks whether removal was successful.',
    'System.out.println("Product deleted successfully.");',
    'This prints a success message.',
    'public void displayInventory() {',
    'This method displays all products currently stored in the inventory.',
    'if (inventory.isEmpty()) {',
    'This checks whether the inventory has any products.',
    'System.out.println("Inventory is empty.");',
    'This prints a message if there are no products.',
    'for (Product1 product : inventory.values()) {',
    'This loops through all products stored in the HashMap.',
    'product.display();',
    'This calls the display method of each Product1 object to print the product details.'
]
for text in items:
    doc.add_paragraph(text)

sec = doc.add_heading('3. Product1.java', level=2)
par = doc.add_paragraph(); par.add_run('Purpose').bold = True
par = doc.add_paragraph('This file represents a single product in the inventory system.')

items = [
    'public class Product1 {',
    'This declares the Product1 class, which acts as a blueprint for each product object.',
    'int productId;',
    'This stores the unique ID of the product.',
    'String productName;',
    'This stores the name of the product.',
    'int quantity;',
    'This stores how many units of the product are available.',
    'double price;',
    'This stores the product price.',
    'public Product1(int productId, String productName, int quantity, double price) {',
    'This is the constructor. It initializes the object with values when a new product is created.',
    'this.productId = productId;',
    'This assigns the passed product ID to the object field.',
    'this.productName = productName;',
    'This assigns the passed product name to the object field.',
    'this.quantity = quantity;',
    'This assigns the passed quantity to the object field.',
    'this.price = price;',
    'This assigns the passed price to the object field.',
    'public void display() {',
    'This method prints the details of a single product.',
    'System.out.println("Product ID   : " + productId);',
    'Prints the product ID.',
    'System.out.println("Product Name : " + productName);',
    'Prints the product name.',
    'System.out.println("Quantity     : " + quantity);',
    'Prints the quantity.',
    'System.out.println("Price        : Rs. " + price);',
    'Prints the price.'
]
for text in items:
    doc.add_paragraph(text)

sec = doc.add_heading('4. Overall Program Flow', level=2)
for text in [
    '1. The program creates an Inventory object.',
    '2. It creates several Product1 objects.',
    '3. It adds those products to the inventory.',
    '4. It displays the inventory list.',
    '5. It updates one product.',
    '6. It deletes another product.',
    '7. It displays the inventory again to show the final result.'
]:
    doc.add_paragraph(text, style='List Bullet')

sec = doc.add_heading('5. Important Java Concepts Used', level=2)
for text in [
    'Class: A blueprint for creating objects.',
    'Object: An instance of a class.',
    'Method: A block of code that performs a task.',
    'Constructor: A special method used to initialize objects.',
    'HashMap: A data structure that stores data as key-value pairs.',
    'Control flow: if-else statements and loops are used to manage logic.'
]:
    doc.add_paragraph(text, style='List Bullet')

sec = doc.add_heading('6. Conclusion', level=2)
par = doc.add_paragraph('This inventory management system is a basic yet practical example of how Java can be used to manage data using classes, objects, and collections. It demonstrates how to add, update, delete, and display records in a simple and organized way.')

doc.save(out)
print(f'Saved: {out}')
